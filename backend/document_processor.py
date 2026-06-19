from typing import List
from io import BytesIO
import os
import tempfile

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader
from docx import Document as DocxDocument

from backend.exceptions import EmptyDocumentError, TextProcessingError


class DocumentProcessor:

    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 100):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            is_separator_regex=False
        )

    def process_file(self, file_bytes: bytes, filename: str) -> List[Document]:
        raw_text = self._extract_text(file_bytes, filename)
        return self._split_into_chunks(raw_text, filename)

    def _split_into_chunks(self, text: str, filename: str) -> List[Document]:
        try:
            metadata = {"source": filename, "type": "user_upload"}

            text_chunks = self.text_splitter.split_text(text)

            documents = [
                Document(page_content=chunk, metadata=metadata)
                for chunk in text_chunks
            ]

            return documents

        except Exception:
            raise TextProcessingError(
                filename=filename,
                chunking_strategy=f"RecursiveCharacter(size={self.chunk_size})"
            )

    def _extract_text(self, file_bytes: bytes, filename: str) -> str:
        ext = os.path.splitext(filename)[1].lower()

        try:
            if ext in [".txt", ".md", ".csv", ".json"]:
                text = self._extract_text_from_plain_text(file_bytes, filename)
            elif ext == ".pdf":
                text = self._extract_text_from_pdf(file_bytes, filename)
            elif ext == ".docx":
                text = self._extract_text_from_docx(file_bytes, filename)
            else:
                raise TextProcessingError(
                    filename=filename,
                    chunking_strategy=f"Unsupported extension: {ext}"
                )
        except EmptyDocumentError:
            raise
        except TextProcessingError:
            raise
        except Exception:
            raise TextProcessingError(
                filename=filename,
                chunking_strategy=f"Extraction for {ext}"
            )

        text = " ".join(text.split())

        if not text.strip():
            raise EmptyDocumentError(filename=filename, file_size=len(text))

        return text

    @staticmethod
    def _extract_text_from_plain_text(file_bytes: bytes, filename: str) -> str:
        try:
            return file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            try:
                return file_bytes.decode("utf-8-sig")
            except UnicodeDecodeError:
                raise TextProcessingError(
                    filename=filename,
                    chunking_strategy="Decode UTF-8"
                )

    @staticmethod
    def _extract_text_from_pdf(file_bytes: bytes, filename: str) -> str:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
            temp_file.write(file_bytes)
            temp_path = temp_file.name

        try:
            loader = PyPDFLoader(temp_path)
            docs = loader.load()
            return "\n".join(doc.page_content for doc in docs)
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)

    @staticmethod
    def _extract_text_from_docx(file_bytes: bytes, filename: str) -> str:
        try:
            doc = DocxDocument(BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n".join(paragraphs)
        except Exception:
            raise TextProcessingError(
                filename=filename,
                chunking_strategy="DOCX extraction"
            )